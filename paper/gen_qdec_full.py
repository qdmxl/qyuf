#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《基于量子酉变换干涉的决策架构》完整版 DOCX。
融合最新 md 文本（43KB 翔实论述）+ 13 张表格 + 6 张图。
输出：量子酉干涉决策架构_final全文版.docx
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = '/home/lijinhan/MXL/科研/ylyw/QYUF/paper'
MD_PATH = os.path.join(BASE, '量子酉干涉决策架构.md')
DOCX_OUT = os.path.join(BASE, '量子酉干涉决策架构_final全文版.docx')
FIG_DIR = os.path.join(BASE, 'figures')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17); s.right_margin = Cm(3.17)

TN = [0]; FN = [0]; ref_section = [False]

def set_font(run, name='Times New Roman', size=11, bold=False, cn='宋体'):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    rpr = run.element.rPr
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        rpr = r.element.rPr
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '黑体')

def add_p(text, indent=True):
    if ref_section[0]:
        pp = doc.add_paragraph()
        r = pp.add_run(text)
        set_font(r, size=10)
        pp.paragraph_format.line_spacing = Pt(16)
        pp.paragraph_format.space_after = Pt(2)
        pp.paragraph_format.first_line_indent = Pt(0)
        return pp
    pp = doc.add_paragraph()
    if indent:
        pp.paragraph_format.first_line_indent = Pt(22)
    pp.paragraph_format.line_spacing = Pt(19)
    pp.paragraph_format.space_after = Pt(4)
    # bold markers ** **
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for pt in parts:
        if pt.startswith('**') and pt.endswith('**'):
            r = pp.add_run(pt[2:-2]); set_font(r, bold=True)
        elif pt.strip():
            r = pp.add_run(pt); set_font(r)
    return pp

def add_center(text, size=11, bold=False, cn='宋体'):
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(text); set_font(r, size=size, bold=bold, cn=cn)
    return pp

def add_table(headers, rows, caption=None):
    if caption:
        TN[0] += 1
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(f'表{TN[0]}  {caption}')
        set_font(r, size=9, bold=True, cn='黑体')
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h); set_font(r, size=9, bold=True, cn='黑体')
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(v)); set_font(r, size=9, cn='宋体')
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_img(fname, caption=None, w=12.5):
    fp = os.path.join(FIG_DIR, fname)
    if not os.path.exists(fp):
        add_p(f'[缺失] {fname}')
        return
    FN[0] += 1
    if caption:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(f'图{FN[0]}  {caption}')
        set_font(r, size=9, bold=True, cn='黑体')
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(); r.add_picture(fp, width=Cm(w))
    doc.add_paragraph()

# =====================================================================
# 读取 md：逐行解析标题/段落；表与图按节插入
# =====================================================================
with open(MD_PATH, 'r') as f:
    md = f.read()

# ---- 表格数据（13 张，取自验证实验）----
T = {}
T[1] = (['场景','Concurrence','rank-1比','结论'],
        [['基态','0','1.000','可分离'],['叠加态','0','1.000','可分离'],
         ['干涉后','0.12–0.18','0.991–0.996','微弱耦合'],['Bell态(对照)','1.00','0.500','最大纠缠']],
        '张量积扩展的Concurrence验证')
T[2] = (['规律','指标','结果','强度','意义'],
        [['单因子翻转','Pearson r','0.66–0.82','强','邻域聚类'],
         ['位序对称','对称度','>0.999','架构保证','U⊗U'],
         ['互补翻转','最强候选对','0.07–0.08','弱','H未编码'],
         ['因子关联','Pearson r','-0.34~0','弱','未编码']],
        '单因子翻转传递性与决策空间邻域结构')
T[3] = (['模式','易理评分','干涉前','干涉后','增益'],
        [['A(最优)','+4','12.5%','44.4%','3.55×'],
         ['B','+2.5','12.5%','21.1%','1.69×'],
         ['C','+1.5','12.5%','15.7%','1.26×'],
         ['D','-4','12.5%','7.5%','0.60×'],
         ['E','0','12.5%','4.2%','0.34×'],
         ['F','0','12.5%','4.1%','0.33×'],
         ['G','-2.5','12.5%','2.1%','0.17×'],
         ['H(最差)','-1.5','12.5%','1.0%','0.08×']],
        '均匀输入下干涉前后概率分布与增益')
T[4] = (['配置','J_adj','J_ying','h_dang','KL散度','Top-5和'],
        [['A(强耦合)','1.5','0.3','0.20','0.30','0.216'],
         ['B(强感应)','0.5','0.8','0.10','0.10','0.172'],
         ['C(平衡)','1.0','0.5','0.30','1.03','0.526']],
        '三种参数配置的主要涌现指标')
T[5] = (['指标','随机 μ','σ','A (ES)','B (ES)','C (ES)'],
        [['KL散度','0.416','0.068','0.304 (-1.6)','0.097 (-4.7)','1.033 (+9.1)'],
         ['Gini','0.493','0.037','0.424 (-1.9)','0.229 (-7.1)','0.706 (+5.8)'],
         ['Top-5和','0.272','0.034','0.216 (-1.6)','0.172 (-2.9)','0.526 (+7.5)'],
         ['对称对','0.697','0.716','1.0 (+0.4)','4.0 (+4.6)','4.0 (+4.6)'],
         ['反对称对','0.633','0.707','0.0 (-0.9)','2.0 (+1.9)','4.0 (+4.8)']],
        'N=1000 Haar随机酉变换统计假设检验')
T[6] = (['τ','平均相关 r','优差比','最佳用例 r'],
        [['0.01','0.053','2.78','-'],['1.00','0.174','2.78','0.313'],
         ['1.53','0.140','-','-'],['2.44','0.128','-','-']],
        'τ参数扫描：决策准确度随判时机的非单调变化')
T[7] = (['方法','平均相关','优差比','提升'],
        [['固定 τ=1.0','0.011','2.78','基线'],['自适应 τ','0.150','3.57','↑12倍']],
        '自适应τ vs 固定τ')
T[8] = (['特征','+方向','-方向','子空间'],
        [['f₀ 硬度','乾 刚','坤 柔','外部'],['f₁ 粗糙度','震 动','巽 入','内部'],
         ['f₂ 规整度','坎 险','离 丽','内部'],['f₃ 动态性','兑 悦','艮 止','外部'],
         ['f₄ 重量','乾 重','坤 轻','外部'],['f₅ 纹理','震 粗','巽 细','内部']],
        '六维物理特征 → 八卦叠加态映射方向')
T[9] = (['物体','初始好模式','涌现好模式','增益','涌现策略'],
        [['海绵','17.6%','50.7%','2.88×','soft_grasp'],
         ['纸团','16.1%','52.7%','3.27×','soft_grasp'],
         ['沙袋','20.7%','62.5%','3.02×','compliant_grasp'],
         ['金属块','98.2%','82.8%','0.84×','power_grasp'],
         ['羽毛','72.2%','72.2%','1.00×','soft_grasp']],
        '五种代表性物体的涌现效果')
T[10] = (['指标','结果'],
         [['总任务','134'],['成功','132'],['成功率','98.5%'],
          ['平均步数','13.0 步'],['总用时','460 s']],
         'ALFWorld V20 valid_unseen 134任务端到端结果')
T[11] = (['涌现模式','次数','占比','成功率'],
         [['主导（precision_grasp）','113','84.3%','98%'],
          ['精密（+6/64）','7','5.2%','100%'],
          ['轻柔（+4/64）','6','4.5%','100%'],
          ['自适应（+3/64）','5','3.7%','100%'],
          ['柔顺（+2/64）','3','2.2%','100%']],
        '涌现决策模式分布')
T[12] = (['方法','成功率','泛化能力'],
         [['V18 专家字典基线','75% (105/140)','仅覆盖21类，3个瓶颈任务失败'],
          ['纯酉干涉+分形','76.4% (107/140)','零样本，自动生成语义签名'],
          ['酉干涉+字典混合','77.2%','与纯酉干涉统计等价']],
        '零样本泛化与字典等价验证 (V18 140局)')
T[13] = (['维度','经典范式','酉干涉涌现范式'],
         [['核心操作','搜索 / 排序 / 匹配','酉变换 / 干涉 / 涌现'],
          ['泛化方式','数据驱动 / 参数搜索','零样本（H谱唯一确定）'],
          ['可解释性','黑箱 + 事后解释','H谱分解可追踪'],
          ['计算复杂度','O(N) 或更差','O(1)（固定维度）'],
          ['规则表示','硬编码 if-then','H 弹性生效']],
        '经典决策范式与酉干涉涌现范式的对比')

# ---- 插图插入位置（按节）----
# 节标题 -> (图文件, 图注)
FIG_POS = {
    ('2.5 时间参数τ：决策时机的数学形式化'): ('tau_scan_yijing_consistency.png',
        'τ从0.01到4.0的扫描：决策准确度呈倒U型，在τ≈0.8–1.5达到峰值'),
    ('2.2 张量积扩展：从8到64'): ('fig3_trigram_joint.png',
        '上下卦联合概率分布 (8×8 八卦空间)：下卦与上卦各覆盖全部8个八卦'),
    ('3.3 干涉效果验证（均匀输入）'): ('fig2_interference.png',
        '均匀输入下酉变换干涉的概率分布重塑'),
    ('3.4 统计假设检验：结构与随机的区别'): ('fig8_statistical_test.png',
        'N=1000 Haar随机对照的统计假设检验结果'),
    ('3.5 τ参数扫描与自适应τ'): ('tau_opt_heatmap.png',
        'τ优化热力图：自适应τ在不同输入下的最优取值'),
    ('4.3 离线物体验证'): ('fig2_final_distribution.png',
        '64维最终涌现概率分布'),
}

# 节标题 -> 要插入的表格编号列表（表号按正文顺序）
TABLE_POS = {
    '3.1 张量积扩展验证': [1],
    '3.2 单因子翻转传递性': [2],
    '3.3 干涉效果验证（均匀输入）': [3],
    '3.4 统计假设检验：结构与随机的区别': [4, 5],
    '3.5 τ参数扫描与自适应τ': [6, 7],
    '4.2 多特征分形编码与策略映射': [8],
    '4.3 离线物体验证': [9],
    '4.4 ALFWorld端到端验证': [10, 11],
    '4.5 零样本泛化与字典等价验证': [12],
    '5.1 与经典决策范式的范式性对比': [13],
}

# =====================================================================
# 渲染标题页 + 摘要 + 关键词
# =====================================================================
def render_header(hbuf):
    text = '\n'.join(hbuf)
    # 主标题
    m = re.search(r'^#\s+(.*)', text, re.M)
    title_main = m.group(1).strip() if m else ''
    # 副标题
    m = re.search(r'^##\s+(.*)', text, re.M)
    title_sub = m.group(1).strip() if m else ''
    # 作者 / 单位行
    authors = []
    for pat in [r'^马兴录.*$', r'^¹.*$', r'^².*$']:
        for mm in re.finditer(pat, text, re.M):
            authors.append(mm.group(0).strip())

    # 标题页
    for _ in range(3): doc.add_paragraph()
    if title_main:
        add_center(title_main, size=22, bold=True, cn='黑体')
    if title_sub:
        doc.add_paragraph()
        add_center(title_sub, size=14, cn='楷体')
    doc.add_paragraph()
    doc.add_paragraph()
    for a in authors:
        add_center(a, size=12, cn='楷体')
    doc.add_page_break()

    # 摘要
    add_h('摘  要', 1)
    m = re.search(r'\*\*摘要\*\*\s*[：:](.*?)(?=\*\*关键词)', text, re.DOTALL)
    if m:
        add_p(m.group(1).strip())
    else:
        # 退而求其次：抓取摘要整段
        m = re.search(r'\*\*摘要\*\*\s*[：:](.*)', text, re.DOTALL)
        if m:
            add_p(m.group(1).strip())
    # 关键词
    km = re.search(r'\*\*关键词\*\*\s*[：:](.*)', text)
    kp = doc.add_paragraph()
    kp.paragraph_format.space_before = Pt(6)
    r = kp.add_run('关键词：'); set_font(r, size=11, bold=True, cn='黑体')
    if km:
        r = kp.add_run(km.group(1).strip()); set_font(r, size=11, cn='楷体')
    doc.add_page_break()


# =====================================================================
# 逐行构建
# =====================================================================
def flush_para(buf):
    text = ' '.join(buf).strip()
    if text:
        add_p(text)

lines = md.split('\n')
buf = []
skip_header = True
header_buf = []          # 暂存头部（标题/作者/摘要）
current_h = None        # 最近渲染的节标题（用于匹配插图）

for line in lines:
    s = line.strip()

    # 头部（--- 之前）：收集标题/作者/摘要/关键词
    if skip_header:
        if s == '---':
            skip_header = False
            render_header(header_buf)   # 渲染标题页+摘要+关键词
        else:
            if s:
                header_buf.append(line)
        continue

    # 参考文献段
    if s.startswith('## ') and '参考文献' in s:
        flush_para(buf); buf = []
        ref_section[0] = True
        add_h('参考文献', 1)
        continue
    if ref_section[0]:
        if s.startswith('[') and ']' in s:
            add_p(s)
        continue

    # 标题
    hm = re.match(r'^(#{1,4})\s+(.*)', line)
    if hm:
        flush_para(buf); buf = []
        level = len(hm.group(1))
        title = hm.group(2).strip()
        add_h(title, level=level)
        current_h = title
        # 在标题后立即插入对应插图与表格
        for key, (fn, cap) in FIG_POS.items():
            if key == title:
                add_img(fn, cap)
        if title in TABLE_POS:
            for tid in TABLE_POS[title]:
                hdr, rows, cap = T[tid]
                add_table(hdr, rows, cap)
        continue

    # 空行 => 段break
    if not s:
        flush_para(buf); buf = []
        continue

    # 其他普通行
    # 表格引用提示（正文里"如表N所示"无需处理，表已按节插入）
    buf.append(s)

flush_para(buf)

doc.save(DOCX_OUT)
print(f'已生成: {DOCX_OUT}')
print(f'大小: {os.path.getsize(DOCX_OUT)/1024:.1f} KB, 表 {TN[0]} 张, 图 {FN[0]} 张')
